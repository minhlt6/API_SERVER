import logging
import os
import re
import uuid
from datetime import datetime, timezone
from typing import List, Optional

from docx import Document as DocxDocument
from fastapi.concurrency import run_in_threadpool
from pypdf import PdfReader
from qdrant_client import QdrantClient
from qdrant_client.http.exceptions import UnexpectedResponse
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    MatchValue,
    PayloadSchemaType,
    PointStruct,
    VectorParams,
)

from .config import CHUNK_OVERLAP, CHUNK_SIZE, QDRANT_API_KEY, QDRANT_COLLECTION, QDRANT_URL
from .document_db import Document, DocumentChunk, SessionLocal
from .models import embeddings

logger = logging.getLogger(__name__)

_ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
_WHITESPACE_RE = re.compile(r"\s+")
_TOKEN_RE = re.compile(r"\S+")


def normalize_text(text: str) -> str:
    if not text:
        return ""

    cleaned = text.replace("\x00", " ")
    cleaned = cleaned.replace("\ufeff", " ")
    cleaned = cleaned.replace("\u200b", " ").replace("\u200c", " ").replace("\u200d", " ")
    cleaned = _WHITESPACE_RE.sub(" ", cleaned)
    return cleaned.strip()


def read_document_content(path: str, extension: str) -> str:
    extension = extension.lower()
    if extension not in _ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported file extension: {extension}")

    if extension == ".pdf":
        reader = PdfReader(path)
        page_texts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(page_texts)

    if extension == ".docx":
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text]

        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    paragraphs.append(" | ".join(row_cells))

        return "\n".join(paragraphs)

    with open(path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()


def chunk_text_by_tokens(text: str, chunk_size: int, overlap: int) -> List[str]:
    if chunk_size <= 0:
        raise ValueError("CHUNK_SIZE must be > 0")
    if overlap < 0:
        raise ValueError("CHUNK_OVERLAP must be >= 0")
    if overlap >= chunk_size:
        raise ValueError("CHUNK_OVERLAP must be smaller than CHUNK_SIZE")

    tokens = _TOKEN_RE.findall(text)
    if not tokens:
        return []

    step = chunk_size - overlap
    chunks: List[str] = []

    for start in range(0, len(tokens), step):
        end = min(start + chunk_size, len(tokens))
        piece = " ".join(tokens[start:end]).strip()
        if piece:
            chunks.append(piece)
        if end >= len(tokens):
            break

    return chunks


def _parse_datetime(value: Optional[str]):
    raw = (value or "").strip()
    if not raw:
        return None

    normalized = raw.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(normalized)
    except ValueError:
        return None


def _ensure_qdrant_collection(client: QdrantClient, vector_size: int, collection_name: str) -> None:
    if not client.collection_exists(collection_name=collection_name):
        client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE),
        )

    _ensure_payload_indexes(client, collection_name)


def _ensure_payload_indexes(client: QdrantClient, collection_name: str) -> None:
    for field_name in ("object_path", "document_id"):
        client.create_payload_index(
            collection_name=collection_name,
            field_name=field_name,
            field_schema=PayloadSchemaType.KEYWORD,
            wait=True,
        )


def _is_missing_payload_index_error(error: Exception) -> bool:
    message = str(error)
    return "Index required but not found" in message


def _delete_existing_document_points(
    client: QdrantClient,
    collection_name: str,
    object_path: Optional[str],
    document_id: str,
) -> None:
    if object_path:
        point_filter = Filter(
            must=[
                FieldCondition(
                    key="object_path",
                    match=MatchValue(value=object_path),
                )
            ]
        )
    else:
        point_filter = Filter(
            must=[
                FieldCondition(
                    key="document_id",
                    match=MatchValue(value=document_id),
                )
            ]
        )

    try:
        client.delete(
            collection_name=collection_name,
            points_selector=point_filter,
            wait=True,
        )
    except UnexpectedResponse as error:
        if not _is_missing_payload_index_error(error):
            raise

        logger.warning(
            "Missing payload index detected while deleting old points in collection=%s. Rebuilding indexes and retrying once.",
            collection_name,
        )
        _ensure_payload_indexes(client, collection_name)
        client.delete(
            collection_name=collection_name,
            points_selector=point_filter,
            wait=True,
        )


def process_document_ingest(
    document_id: str,
    file_path: Optional[str] = None,
    collection_name: Optional[str] = None,
    source_path: Optional[str] = None,
    source_object_path: Optional[str] = None,
    source_updated_at: Optional[str] = None,
    source_etag: Optional[str] = None,
    cleanup_file: bool = False,
    size: Optional[int] = None,
) -> bool:
    db = SessionLocal()

    effective_file_path = (file_path or "").strip()
    effective_source_path = (source_path or "").strip()

    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        if document is None:
            logger.error("Document not found for ingest: %s", document_id)
            return False

        document.status = "processing"
        document.error_message = None
        db.commit()

        ingest_file_path = effective_file_path or document.path
        if not ingest_file_path:
            raise ValueError("Document file path is missing for ingest.")

        source_object_ref = (source_object_path or document.object_path or "").strip() or None

        extension_source = source_object_ref or document.stored_name or ingest_file_path
        _, extension = os.path.splitext(extension_source)

        raw_text = read_document_content(ingest_file_path, extension)
        normalized = normalize_text(raw_text)
        chunks = chunk_text_by_tokens(normalized, CHUNK_SIZE, CHUNK_OVERLAP)

        if not chunks:
            raise ValueError("Document has no readable content after normalization.")

        if not QDRANT_URL:
            raise ValueError("QDRANT_URL is required for ingest.")

        client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY or None)
        vectors = embeddings.embed_documents(chunks)

        if not vectors or not vectors[0]:
            raise ValueError("Failed to create embeddings for chunks.")

        target_collection = (collection_name or document.collection_name or QDRANT_COLLECTION or "").strip()
        if not target_collection:
            raise ValueError("Target collection is empty.")

        _ensure_qdrant_collection(client, len(vectors[0]), target_collection)
        _delete_existing_document_points(client, target_collection, source_object_ref, document.id)

        created_at = datetime.now(timezone.utc).isoformat()
        points: List[PointStruct] = []
        db_chunk_rows: List[DocumentChunk] = []

        for index, (chunk_text, vector) in enumerate(zip(chunks, vectors)):
            point_id = str(uuid.uuid4())
            payload = {
                "document_id": document.id,
                "filename": document.original_name,
                "stored_name": document.stored_name,
                "path": effective_source_path or document.path,
                "object_path": source_object_ref,
                "folder_key": document.folder_key,
                "collection_name": target_collection,
                "source_updated_at": source_updated_at,
                "source_etag": source_etag,
                "chunk_index": index,
                "created_at": created_at,
                "content": chunk_text,
            }

            points.append(PointStruct(id=point_id, vector=vector, payload=payload))
            db_chunk_rows.append(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=index,
                    content_preview=chunk_text[:200],
                    qdrant_point_id=point_id,
                )
            )

        client.upsert(collection_name=target_collection, points=points, wait=True)

        db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()
        db.bulk_save_objects(db_chunk_rows)

        if effective_source_path:
            document.path = effective_source_path
        if source_object_ref:
            document.object_path = source_object_ref
        if source_etag:
            document.source_etag = source_etag
        if source_updated_at:
            parsed_source_updated = _parse_datetime(source_updated_at)
            if parsed_source_updated is not None:
                document.source_updated_at = parsed_source_updated
        if size is not None:
            document.size = int(size)

        document.collection_name = target_collection
        document.last_synced_at = datetime.now(timezone.utc)
        document.deleted_at = None
        document.total_chunks = len(chunks)
        document.status = "done"
        document.error_message = None
        db.commit()

        logger.info("Document ingest success. document_id=%s total_chunks=%s", document.id, len(chunks))
        return True
    except Exception as error:
        db.rollback()

        failed_doc = db.query(Document).filter(Document.id == document_id).first()
        if failed_doc is not None:
            failed_doc.status = "failed"
            failed_doc.error_message = str(error)
            db.commit()

        logger.exception("Document ingest failed. document_id=%s", document_id)
        return False
    finally:
        if cleanup_file and effective_file_path and os.path.exists(effective_file_path):
            try:
                os.remove(effective_file_path)
            except Exception:
                logger.exception("Failed to remove temporary ingest file: %s", effective_file_path)

        db.close()


def delete_vectors_for_object_path(collection_name: str, object_path: str) -> bool:
    if not QDRANT_URL:
        return False

    target_collection = (collection_name or "").strip()
    normalized_object_path = (object_path or "").strip()
    if not target_collection or not normalized_object_path:
        return False

    client = QdrantClient(url=QDRANT_URL, api_key=QDRANT_API_KEY or None)

    if not client.collection_exists(collection_name=target_collection):
        return False

    point_filter = Filter(
        must=[
            FieldCondition(
                key="object_path",
                match=MatchValue(value=normalized_object_path),
            )
        ]
    )

    try:
        _ensure_payload_indexes(client, target_collection)
        client.delete(
            collection_name=target_collection,
            points_selector=point_filter,
            wait=True,
        )
    except UnexpectedResponse as error:
        if not _is_missing_payload_index_error(error):
            raise

        logger.warning(
            "Missing payload index detected while deleting object_path in collection=%s. Rebuilding indexes and retrying once.",
            target_collection,
        )
        _ensure_payload_indexes(client, target_collection)
        client.delete(
            collection_name=target_collection,
            points_selector=point_filter,
            wait=True,
        )

    return True


async def run_document_ingest_task(document_id: str) -> None:
    # Heavy ingest work runs in threadpool to keep event loop responsive.
    await run_in_threadpool(process_document_ingest, document_id)
