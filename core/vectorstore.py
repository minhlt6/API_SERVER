import asyncio
import logging
import os
import re
from typing import Any, Dict, List

import pdfplumber
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document as LangChainDocument

from .text_utils import clean_text

logger = logging.getLogger(__name__)

ACADEMIC_YEAR_PATTERN = re.compile(r"(20\d{2})\s*[-_]\s*(20\d{2})")


def normalize_academic_year(start_year: str, end_year: str) -> str:
    return f"{int(start_year):04d}-{int(end_year):04d}"


def extract_academic_year(text: str) -> str:
    if not text:
        return ""
    match = ACADEMIC_YEAR_PATTERN.search(text)
    if not match:
        return ""
    return normalize_academic_year(match.group(1), match.group(2))


def table_to_markdown(data: List[List[str]]) -> str:
    if not data or len(data) < 2:
        return ""

    header = [str(cell).replace("\n", " ").strip() if cell else "" for cell in data[0]]
    separator = ["---"] * len(header)
    markdown_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]

    for row in data[1:]:
        clean_row = [str(cell).replace("\n", "<br>").strip() if cell else "" for cell in row]
        markdown_lines.append("| " + " | ".join(clean_row) + " |")

    return "\n".join(markdown_lines) + "\n\n"


def read_pdf_with_tables(filepath: str) -> List[LangChainDocument]:
    docs: List[LangChainDocument] = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_index, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                table_texts: List[str] = []
                if tables:
                    for table in tables:
                        md_table = table_to_markdown(table)
                        if md_table:
                            table_texts.append(md_table)

                full_content = text + "\n\n[BANG DU LIEU TRICH XUAT]:\n" + "\n".join(table_texts)
                if full_content.strip():
                    docs.append(
                        LangChainDocument(
                            page_content=full_content,
                            metadata={"source": filepath, "page": page_index},
                        )
                    )
    except Exception as error:
        logger.error("Loi doc PDF (pdfplumber) %s: %s", os.path.basename(filepath), error)

    return docs


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Chỉ hỗ trợ duyệt Document hoặc Cell")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def read_docx_with_tables(filepath: str) -> str:
    doc = Document(filepath)
    full_text: List[str] = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.strip():
                full_text.append(block.text.strip())
        elif isinstance(block, Table):
            table_data: List[List[str]] = []
            for row in block.rows:
                row_data: List[str] = []
                for cell in row.cells:
                    row_data.append(clean_text(cell.text))
                table_data.append(row_data)

            md_table = table_to_markdown(table_data)
            if md_table:
                full_text.append(f"\n{md_table}\n")

    return "\n".join(full_text)


def load_documents_from_file(filepath: str, filename: str) -> List[LangChainDocument]:
    docs: List[LangChainDocument] = []
    lower_name = filename.lower()

    try:
        if lower_name.endswith(".pdf"):
            docs = read_pdf_with_tables(filepath)
        elif lower_name.endswith(".docx"):
            text = read_docx_with_tables(filepath)
            if text:
                docs = [LangChainDocument(page_content=text, metadata={"source": filepath})]
        elif lower_name.endswith(".txt"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as input_file:
                text = input_file.read()
            if text and text.strip():
                docs = [LangChainDocument(page_content=text, metadata={"source": filepath})]

        if docs:
            logger.info("Da doc: %s", filename)

        return docs
    except Exception as error:
        logger.error("Loi doc %s: %s", filename, str(error)[:120])
        return []


async def build_vectorstore_improved(
    sync_coordinator: Any,
    startup_wait_seconds: int = 5,
) -> Dict[str, Any]:
    """Supabase build step: trigger one initial sync and optionally wait for completion."""
    if sync_coordinator is None:
        raise ValueError("sync_coordinator is required")

    startup_sync_task = asyncio.create_task(
        sync_coordinator.run_sync(
            trigger="startup:initial_sync",
            queue_if_locked=False,
        )
    )

    if startup_wait_seconds <= 0:
        return {
            "task": startup_sync_task,
            "initial_sync": None,
            "timed_out": True,
        }

    try:
        initial_sync = await asyncio.wait_for(
            asyncio.shield(startup_sync_task),
            timeout=startup_wait_seconds,
        )
        return {
            "task": startup_sync_task,
            "initial_sync": initial_sync,
            "timed_out": False,
        }
    except asyncio.TimeoutError:
        return {
            "task": startup_sync_task,
            "initial_sync": None,
            "timed_out": True,
        }


def load_vectorstore_improved(sync_coordinator: Any) -> Dict[str, Any]:
    """Supabase load step: return current coordinator health snapshot."""
    if sync_coordinator is None:
        return {}

    try:
        state = sync_coordinator.get_health_snapshot()
        return state if isinstance(state, dict) else {}
    except Exception:
        logger.exception("Khong the lay sync state tu coordinator")
        return {}