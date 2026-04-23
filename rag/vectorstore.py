import asyncio
import logging
import os
from typing import Any, Dict, List

import pdfplumber
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document as LangChainDocument

from utils.text_utils import clean_text

logger = logging.getLogger(__name__)

def table_to_unrolled_text(data: List[List[str]], is_docx: bool = False) -> str:
    
    if not data or len(data) < 2:
        return ""

    # Làm sạch dữ liệu ban đầu chuyển None thành chuỗi rỗng
    cleaned_data = []
    for row in data:
        cleaned_row = [str(cell).strip() if cell else "" for cell in row]
        cleaned_data.append(cleaned_row)

    num_cols = len(cleaned_data[0])
    header_row = cleaned_data[0]

    # CHỈ CHẠY FORWARD FILL NẾU KHÔNG PHẢI FILE WORD
    if not is_docx:
        # 2. Kỹ thuật Forward-Fill cho khu vực Header (Xử lý gộp cột - Colspan)
        # Giả định hàng đầu tiên chắc chắn là Header
        for i in range(1, num_cols):
            if not header_row[i] and header_row[i-1]:
                header_row[i] = header_row[i-1] # Kéo giá trị từ trái sang phải

        # 3. Kỹ thuật Forward-Fill cho khu vực Dữ liệu (Xử lý gộp hàng - Rowspan)
        for r in range(1, len(cleaned_data)):
            for c in range(num_cols):
                # Nếu ô hiện tại rỗng, kéo giá trị từ ô ngay bên trên xuống
                if not cleaned_data[r][c] and cleaned_data[r-1][c]:
                    cleaned_data[r][c] = cleaned_data[r-1][c]

    # 4. Trải phẳng bảng (Unrolling)
    headers = cleaned_data[0]
    unrolled_rows = []

    for r in range(1, len(cleaned_data)):
        row_values = cleaned_data[r]
        row_text_parts = []
        
        # Chỉ ghép những ô có dữ liệu thực sự (khác Header)
        for c in range(min(len(headers), len(row_values))):
            header_val = headers[c]
            cell_val = row_values[c]
            
            # Tránh lặp lại nếu dữ liệu vô tình giống hệt Header
            if cell_val and cell_val != header_val:
                row_text_parts.append(f"{header_val}: {cell_val}")
                
        if row_text_parts:
            unrolled_rows.append("- " + " | ".join(row_text_parts))

    return "\n" + "\n".join(unrolled_rows) + "\n\n"

def read_pdf_with_tables(filepath: str) -> List[LangChainDocument]:
    docs: List[LangChainDocument] = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_index, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                table_texts: List[str] = []
                if tables:
                    for table in tables:
                        # Vẫn chạy Forward-Fill bình thường cho PDF
                        unrolled_table = table_to_unrolled_text(table, is_docx=False)
                        if unrolled_table:
                            table_texts.append(unrolled_table)

                full_content = text + "\n\n[BANG DU LIEU TRICH XUAT]:\n" + "\n".join(table_texts)
                if full_content.strip():
                    docs.append(
                        LangChainDocument(
                            page_content=full_content,
                            metadata={"source": filepath, "page": page_index},
                        )
                    )
    except Exception as error:
        logger.error("Lỗi đọc PDF  %s: %s", os.path.basename(filepath), error)

    return docs

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Chỉ hỗ trợ duyệt Document hoặc Cell")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def read_docx_with_tables(filepath: str) -> str:
    doc = Document(filepath)
    full_text: List[str] = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if block.text.strip():
                full_text.append(block.text.strip())
        elif isinstance(block, Table):
            table_data: List[List[str]] = []
            for row in block.rows:
                row_data: List[str] = []
                for cell in row.cells:
                    row_data.append(clean_text(cell.text))
                table_data.append(row_data)

            # CẮT FORWARD-FILL TẠI ĐÂY BẰNG is_docx=True
            unrolled_table = table_to_unrolled_text(table_data, is_docx=True)
            if unrolled_table:
                full_text.append(f"\n{unrolled_table}\n")

    return "\n".join(full_text)

def load_documents_from_file(filepath: str, filename: str) -> List[LangChainDocument]:
    docs: List[LangChainDocument] = []
    lower_name = filename.lower()

    try:
        if lower_name.endswith(".pdf"):
            docs = read_pdf_with_tables(filepath)
        elif lower_name.endswith(".docx"):
            text = read_docx_with_tables(filepath)
            if text:
                docs = [LangChainDocument(page_content=text, metadata={"source": filepath})]
        elif lower_name.endswith(".txt"):
            with open(filepath, "r", encoding="utf-8", errors="ignore") as input_file:
                text = input_file.read()
            if text and text.strip():
                docs = [LangChainDocument(page_content=text, metadata={"source": filepath})]

        if docs:
            logger.info("Da doc: %s", filename)

        return docs
    except Exception as error:
        logger.error("Loi doc %s: %s", filename, str(error)[:120])
        return []

async def build_vectorstore_improved(
    sync_coordinator: Any,
    startup_wait_seconds: int = 5,
) -> Dict[str, Any]:
    if sync_coordinator is None:
        raise ValueError("sync_coordinator is required")

    startup_sync_task = asyncio.create_task(
        sync_coordinator.run_sync(
            trigger="startup:initial_sync",
            queue_if_locked=False,
        )
    )

    if startup_wait_seconds <= 0:
        return {
            "task": startup_sync_task,
            "initial_sync": None,
            "timed_out": True,
        }

    try:
        initial_sync = await asyncio.wait_for(
            asyncio.shield(startup_sync_task),
            timeout=startup_wait_seconds,
        )
        return {
            "task": startup_sync_task,
            "initial_sync": initial_sync,
            "timed_out": False,
        }
    except asyncio.TimeoutError:
        return {
            "task": startup_sync_task,
            "initial_sync": None,
            "timed_out": True,
        }

def load_vectorstore_improved(sync_coordinator: Any) -> Dict[str, Any]:
    if sync_coordinator is None:
        return {}
    try:
        state = sync_coordinator.get_health_snapshot()
        return state if isinstance(state, dict) else {}
    except Exception:
        logger.exception("Khong the lay sync state tu coordinator")
        return {}